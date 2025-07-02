import re
import spacy
from spacy.tokens import Doc
from typing import List, Dict, Optional
import PyPDF2
import pdfplumber
from pathlib import Path
import io
from docx import Document

class TextPreprocessor:
    """Handles text cleaning and preprocessing for stroke reports."""
    
    def __init__(self):
        # Load German spaCy model
        try:
            self.nlp = spacy.load("de_core_news_sm")
        except OSError:
            print("German spaCy model not found. Please install with:")
            print("python -m spacy download de_core_news_sm")
            raise
    
    def clean_text(self, text: str) -> str:
        """Basic text cleaning."""
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove common headers/footers patterns
        text = re.sub(r'(Befund|Diagnose|Impression):\s*', '', text, flags=re.IGNORECASE)
        text = re.sub(r'(Dr\.|Prof\.|Oberarzt|Chefarzt)\s+\w+', '', text)
        
        # Strip and lowercase
        return text.strip().lower()
    
    def tokenize_and_tag(self, text: str) -> Doc:
        """Process text with spaCy for POS and NER."""
        return self.nlp(text)
    
    def extract_sentences(self, text: str) -> List[str]:
        """Extract individual sentences."""
        doc = self.nlp(text)
        return [sent.text.strip() for sent in doc.sents]
    
    def extract_text_from_pdf(self, pdf_path: str, method: str = "pdfplumber") -> str:
        """
        Extract text from PDF file using multiple methods.
        
        Args:
            pdf_path: Path to PDF file
            method: Extraction method ('pdfplumber' or 'pypdf2')
            
        Returns:
            Extracted text as string
        """
        pdf_path = Path(pdf_path)
        if not pdf_path.exists():
            raise FileNotFoundError(f"PDF file not found: {pdf_path}")
        
        if method == "pdfplumber":
            return self._extract_with_pdfplumber(pdf_path)
        elif method == "pypdf2":
            return self._extract_with_pypdf2(pdf_path)
        else:
            raise ValueError(f"Unknown extraction method: {method}")
    
    def _extract_with_pdfplumber(self, pdf_path: Path) -> str:
        """Extract text using pdfplumber (better for complex layouts)."""
        extracted_text = ""
        
        try:
            with pdfplumber.open(pdf_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n"
                        extracted_text += page_text
                        extracted_text += "\n"
        except Exception as e:
            raise RuntimeError(f"Error extracting PDF with pdfplumber: {e}")
        
        return extracted_text.strip()
    
    def _extract_with_pypdf2(self, pdf_path: Path) -> str:
        """Extract text using PyPDF2 (faster but less accurate)."""
        extracted_text = ""
        
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_text += f"\n--- Page {page_num + 1} ---\n"
                        extracted_text += page_text
                        extracted_text += "\n"
        except Exception as e:
            raise RuntimeError(f"Error extracting PDF with PyPDF2: {e}")
        
        return extracted_text.strip()
    
    def extract_text_from_docx(self, docx_path: str) -> str:
        """
        Extract text from Word document (.docx).
        
        Args:
            docx_path: Path to Word document file
            
        Returns:
            Extracted text as string
        """
        docx_path = Path(docx_path)
        if not docx_path.exists():
            raise FileNotFoundError(f"Word document not found: {docx_path}")
        
        try:
            doc = Document(docx_path)
            extracted_text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            extracted_text += cell.text + " "
                    extracted_text += "\n"
            
            return extracted_text.strip()
            
        except Exception as e:
            raise RuntimeError(f"Error extracting Word document: {e}")
    
    def process_pdf_batch(self, pdf_folder: str, output_folder: str = None) -> List[Dict[str, str]]:
        """
        Process multiple PDF files in a folder.
        
        Args:
            pdf_folder: Folder containing PDF files
            output_folder: Optional folder to save extracted text files
            
        Returns:
            List of dictionaries with filename and extracted text
        """
        pdf_folder = Path(pdf_folder)
        if not pdf_folder.exists():
            raise FileNotFoundError(f"PDF folder not found: {pdf_folder}")
        
        results = []
        pdf_files = list(pdf_folder.glob("*.pdf"))
        
        if not pdf_files:
            print(f"No PDF files found in {pdf_folder}")
            return results
        
        # Create output folder if specified
        if output_folder:
            output_folder = Path(output_folder)
            output_folder.mkdir(parents=True, exist_ok=True)
        
        for pdf_file in pdf_files:
            try:
                print(f"Processing: {pdf_file.name}")
                
                # Try pdfplumber first, fallback to PyPDF2
                try:
                    extracted_text = self.extract_text_from_pdf(pdf_file, method="pdfplumber")
                except Exception:
                    print(f"  pdfplumber failed, trying PyPDF2...")
                    extracted_text = self.extract_text_from_pdf(pdf_file, method="pypdf2")
                
                # Clean the extracted text
                cleaned_text = self.clean_text(extracted_text)
                
                result = {
                    'filename': pdf_file.name,
                    'filepath': str(pdf_file),
                    'raw_text': extracted_text,
                    'cleaned_text': cleaned_text,
                    'char_count': len(cleaned_text),
                    'word_count': len(cleaned_text.split()) if cleaned_text else 0
                }
                
                results.append(result)
                
                # Save to output folder if specified
                if output_folder:
                    output_file = output_folder / f"{pdf_file.stem}.txt"
                    with open(output_file, 'w', encoding='utf-8') as f:
                        f.write(cleaned_text)
                    result['output_file'] = str(output_file)
                
                print(f"  ✅ Extracted {result['word_count']} words")
                
            except Exception as e:
                print(f"  ❌ Error processing {pdf_file.name}: {e}")
                results.append({
                    'filename': pdf_file.name,
                    'filepath': str(pdf_file),
                    'error': str(e),
                    'raw_text': None,
                    'cleaned_text': None
                })
        
        return results
    
    def preprocess_for_training(self, text: str, 
                              remove_patient_info: bool = True,
                              normalize_medical_terms: bool = True) -> str:
        """
        Advanced preprocessing for training data.
        
        Args:
            text: Input text
            remove_patient_info: Remove potential patient identifiers
            normalize_medical_terms: Standardize medical terminology
            
        Returns:
            Processed text suitable for training
        """
        # Start with basic cleaning
        processed_text = self.clean_text(text)
        
        # Remove potential patient identifiers
        if remove_patient_info:
            # Remove ages, dates, names patterns
            processed_text = re.sub(r'\b\d{1,3}[-\s]?jährig[er]?\b', '[AGE]', processed_text)
            processed_text = re.sub(r'\b\d{1,2}[./]\d{1,2}[./]\d{2,4}\b', '[DATE]', processed_text)
            processed_text = re.sub(r'\b[A-Z][a-z]+\s+[A-Z][a-z]+\b', '[NAME]', processed_text)
        
        # Normalize medical terms
        if normalize_medical_terms:
            # Standardize common variations
            medical_normalizations = {
                r'\brt[_\-\s]?pa\b': 'rtpa',
                r'\bt[_\-\s]?pa\b': 'rtpa',
                r'\btici[_\-\s]?(\d[abc]?)\b': r'tici \1',
                r'\bstent[_\-\s]?retriever\b': 'stentretriever',
                r'\bmechanisch[e]?\s+thrombektomie\b': 'mechanische thrombektomie',
                r'\ballgemein[_\-\s]?anästhesie\b': 'allgemeinanästhesie',
                r'\blokal[_\-\s]?anästhesie\b': 'lokalanästhesie'
            }
            
            for pattern, replacement in medical_normalizations.items():
                processed_text = re.sub(pattern, replacement, processed_text, flags=re.IGNORECASE)
        
        return processed_text.strip()
    
    def validate_pdf_extraction(self, pdf_path: str) -> Dict[str, any]:
        """
        Validate PDF extraction quality.
        
        Args:
            pdf_path: Path to PDF file
            
        Returns:
            Validation metrics
        """
        try:
            # Extract with both methods
            text_pdfplumber = self.extract_text_from_pdf(pdf_path, "pdfplumber")
            text_pypdf2 = self.extract_text_from_pdf(pdf_path, "pypdf2")
            
            # Calculate metrics
            metrics = {
                'pdf_file': pdf_path,
                'pdfplumber_chars': len(text_pdfplumber),
                'pypdf2_chars': len(text_pypdf2),
                'pdfplumber_words': len(text_pdfplumber.split()),
                'pypdf2_words': len(text_pypdf2.split()),
                'similarity': self._calculate_text_similarity(text_pdfplumber, text_pypdf2),
                'recommended_method': 'pdfplumber' if len(text_pdfplumber) > len(text_pypdf2) else 'pypdf2',
                'extraction_quality': 'good' if len(text_pdfplumber) > 100 else 'poor'
            }
            
            return metrics
            
        except Exception as e:
            return {
                'pdf_file': pdf_path,
                'error': str(e),
                'extraction_quality': 'failed'
            }
    
    def _calculate_text_similarity(self, text1: str, text2: str) -> float:
        """Calculate basic similarity between two texts."""
        if not text1 or not text2:
            return 0.0
        
        # Simple word-based similarity
        words1 = set(text1.lower().split())
        words2 = set(text2.lower().split())
        
        if not words1 and not words2:
            return 1.0
        
        intersection = words1.intersection(words2)
        union = words1.union(words2)
        
        return len(intersection) / len(union) if union else 0.0